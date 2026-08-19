import sys
import json
from pathlib import Path
import requests

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

# Configuration
INPUT_DIR = Path(r"c:\Users\VINOTHINI B\OneDrive\Desktop\GLM - OCR\apps\input")
OUTPUT_DIR = Path(r"c:\Users\VINOTHINI B\OneDrive\Desktop\GLM - OCR\apps\output_benchmarks")
SDK_SERVER_URL = "http://localhost:5002/glmocr/parse"

BENCHMARK_CHECKPOINTS = [
    "Page loading",
    "Layout detection",
    "GLM-OCR model inference",
    "Post-processing/markdown generation",
]


def scan_pdfs():
    """Scan INPUT_DIR for PDF files."""
    if not INPUT_DIR.exists():
        print(f"Error: Input directory {INPUT_DIR} does not exist.")
        return []
    return sorted(list(INPUT_DIR.glob("*.pdf")), key=lambda p: p.stat().st_size)


def check_checkpoint_coverage(checkpoints):
    """Return ordered checkpoint entries and any missing benchmark stages."""
    checkpoint_map = {entry["checkpoint"]: entry for entry in checkpoints}
    ordered = []
    missing = []
    for checkpoint_name in BENCHMARK_CHECKPOINTS:
        checkpoint = checkpoint_map.get(checkpoint_name)
        if checkpoint is None:
            missing.append(checkpoint_name)
            continue
        ordered.append(checkpoint)
    return ordered, missing


def normalize_benchmark_timings(timings_ms):
    """Convert pipeline timing metadata into the checkpoint format expected by the benchmark report."""
    normalized = []
    mapping = {
        "page_loading_ms": ("Page loading", "Rendered PDF pages with the pipeline page loader"),
        "layout_detection_ms": ("Layout detection", "Detected page regions and structural layout"),
        "ocr_inference_ms": ("GLM-OCR model inference", "Executed OCR inference over detected regions"),
        "postprocessing_ms": ("Post-processing/markdown generation", "Merged OCR output and formatted final markdown"),
    }
    for key, (checkpoint_name, description) in mapping.items():
        duration_ms = float(timings_ms.get(key, 0.0) or 0.0)
        normalized.append(
            {
                "step": len(normalized) + 1,
                "checkpoint": checkpoint_name,
                "duration": duration_ms / 1000.0,
                "description": description,
            }
        )
    return normalized


def export_word_report(benchmark_data, output_path: Path):
    """Export the benchmark timing results to a Microsoft Word document."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx is not installed; Word export skipped.")
        return None

    ordered_checkpoints, missing = check_checkpoint_coverage(benchmark_data.get("checkpoints", []))
    report_path = Path(output_path)
    report_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title = doc.add_paragraph()
    title_run = title.add_run("GLM-OCR")
    title_run.bold = True
    title_run.font.size = 22
    title_run.font.color.rgb = None
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Performance Benchmark Report")
    subtitle_run.bold = True
    subtitle_run.font.size = 26

    doc.add_paragraph(f"Target file: {benchmark_data.get('filename', 'Unknown file')}")
    doc.add_paragraph(f"Pages: {benchmark_data.get('pages', 0)}")
    doc.add_paragraph(f"File size: {benchmark_data.get('size_mb', 0):.2f} MB")
    doc.add_paragraph(f"Total execution time: {benchmark_data.get('total_time', 0):.3f} seconds")

    doc.add_heading("Processing Duration Breakdown", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Step"
    table.rows[0].cells[1].text = "Checkpoint"
    table.rows[0].cells[2].text = "Duration (s)"
    table.rows[0].cells[3].text = "Details"

    for idx, checkpoint in enumerate(ordered_checkpoints, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = checkpoint.get("checkpoint", "Unknown")
        row_cells[2].text = f"{checkpoint.get('duration', 0):.3f}"
        row_cells[3].text = checkpoint.get("description", "")

    doc.add_heading("Checkpoint Coverage", level=1)
    for checkpoint_name in BENCHMARK_CHECKPOINTS:
        checkbox = "[x]" if checkpoint_name not in missing else "[ ]"
        doc.add_paragraph(f"{checkbox} {checkpoint_name}")

    if missing:
        doc.add_paragraph("Missing checkpoints from the current run: " + ", ".join(missing))

    doc.add_heading("Benchmark Performance Summary", level=1)
    doc.add_paragraph(
        "This benchmark captures the end-to-end processing time for the OCR pipeline, "
        "including PDF ingestion, layout analysis, region preparation, and model inference."
    )
    doc.add_paragraph(
        "The benchmark is designed to measure the time spent in each pipeline stage so performance "
        "bottlenecks can be identified and optimized."
    )

    doc.save(report_path)
    return report_path


def parse_pdf_with_benchmarks(pdf_path: Path):
    """Parse a single PDF, measuring checkpoints and saving JSON/MD results."""
    print("\n==============================================")
    print(f"Benchmarking OCR for: {pdf_path.name}")
    print("==============================================")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("Reading PDF metadata for reporting...")
    num_pages = 0
    file_size_mb = pdf_path.stat().st_size / (1024 * 1024)

    if fitz is not None:
        try:
            doc = fitz.open(pdf_path)
            num_pages = len(doc)
            doc.close()
        except Exception as e:
            print(f"Warning: Failed to read PDF metadata via PyMuPDF: {e}")
            num_pages = 1
    else:
        num_pages = 1  # Fallback

    print("Requesting OCR; timings come from internal pipeline instrumentation...")
    file_uri = f"file:///{pdf_path.absolute().as_posix()}"
    payload = {"images": [file_uri]}

    try:
        response = requests.post(SDK_SERVER_URL, json=payload, timeout=60)
        response.raise_for_status()
        resp_data = response.json()
    except Exception as e:
        print(f"Error calling GLM-OCR SDK server: {e}")
        print("Make sure the SDK Flask server is running at http://localhost:5002")
        return None

    checkpoints = normalize_benchmark_timings(resp_data.get("timings_ms", {}))
    total_time = sum(c["duration"] for c in checkpoints)
    print(f"Successfully parsed '{pdf_path.name}' in {total_time:.2f}s!")

    # Extract results
    json_result = resp_data.get("json_result", [])
    markdown_result = resp_data.get("markdown_result", "")

    # Save outputs
    stem = pdf_path.stem
    json_out_file = OUTPUT_DIR / f"{stem}.json"
    md_out_file = OUTPUT_DIR / f"{stem}.md"

    with open(json_out_file, "w", encoding="utf-8") as f:
        json.dump(json_result, f, ensure_ascii=False, indent=2)

    with open(md_out_file, "w", encoding="utf-8") as f:
        f.write(markdown_result)

    print(f"Saved JSON output to: {json_out_file}")
    print(f"Saved MD output to: {md_out_file}")

    return {
        "filename": pdf_path.name,
        "size_mb": file_size_mb,
        "pages": num_pages,
        "checkpoints": checkpoints,
        "total_time": total_time,
        "markdown": markdown_result,
        "json": json_result,
        "json_path": str(json_out_file.absolute()),
        "md_path": str(md_out_file.absolute()),
    }


def generate_dashboard_html(benchmark_data, pdf_list):
    """Generate a gorgeous dark-themed HTML dashboard matching the design guidelines."""
    # Build list of PDFs for the sidebar
    sidebar_items = ""
    for pdf in pdf_list:
        active_class = "active" if pdf.name == benchmark_data["filename"] else ""
        size_mb = pdf.stat().st_size / (1024 * 1024)
        sidebar_items += f"""
        <div class="sidebar-item {active_class}" onclick="window.location.reload()">
            <div class="pdf-icon">📄</div>
            <div class="pdf-info">
                <div class="pdf-name">{pdf.name}</div>
                <div class="pdf-meta">{size_mb:.2f} MB</div>
            </div>
        </div>
        """

    # Build timing table rows
    table_rows = ""
    for c in benchmark_data["checkpoints"]:
        table_rows += f"""
        <tr>
            <td class="step-col">{c['step']}</td>
            <td class="checkpoint-col">{c['checkpoint']}</td>
            <td class="duration-col">{c['duration']:.3f}s</td>
            <td class="desc-col">{c['description']}</td>
        </tr>
        """

    # Build CSS bar charts
    chart_bars = ""
    colors = ["#3b82f6", "#10b981", "#f59e0b", "#ef4444", "#8b5cf6", "#ec4899", "#14b8a6"]
    for idx, c in enumerate(benchmark_data["checkpoints"]):
        pct = (c["duration"] / benchmark_data["total_time"]) * 100
        color = colors[idx % len(colors)]
        chart_bars += f"""
        <div class="chart-item">
            <div class="chart-label">
                <span>{c['checkpoint']}</span>
                <span>{pct:.1f}% ({c['duration']:.2f}s)</span>
            </div>
            <div class="chart-bar-container">
                <div class="chart-bar" style="width: {pct}%; background-color: {color};"></div>
            </div>
        </div>
        """

    # Format JSON safely for HTML pre representation
    json_formatted = json.dumps(benchmark_data["json"], ensure_ascii=False, indent=2)

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>⚡ GLM-OCR Runtime Dashboard</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-color: #0b0f19;
            --card-bg: #151d30;
            --text-primary: #f3f4f6;
            --text-secondary: #9ca3af;
            --accent-blue: #3b82f6;
            --accent-green: #10b981;
            --accent-purple: #8b5cf6;
            --border-color: #243049;
            --shadow: 0 8px 30px rgba(0,0,0,0.4);
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}

        body {{
            font-family: 'Outfit', sans-serif;
            background-color: var(--bg-color);
            color: var(--text-primary);
            display: flex;
            height: 100vh;
            overflow: hidden;
        }}

        /* Sidebar styling */
        .sidebar {{
            width: 320px;
            background-color: #0e1424;
            border-right: 1px solid var(--border-color);
            display: flex;
            flex-direction: column;
            flex-shrink: 0;
        }}

        .sidebar-header {{
            padding: 24px;
            border-bottom: 1px solid var(--border-color);
        }}

        .sidebar-title {{
            font-size: 1.25rem;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 8px;
            color: #ffffff;
        }}

        .sidebar-list {{
            flex: 1;
            overflow-y: auto;
            padding: 16px 12px;
        }}

        .sidebar-item {{
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 12px;
            border-radius: 10px;
            cursor: pointer;
            margin-bottom: 8px;
            transition: all 0.2s ease;
            border: 1px solid transparent;
        }}

        .sidebar-item:hover {{
            background-color: rgba(59, 130, 246, 0.08);
            border-color: rgba(59, 130, 246, 0.2);
        }}

        .sidebar-item.active {{
            background-color: rgba(59, 130, 246, 0.15);
            border-color: var(--accent-blue);
        }}

        .pdf-icon {{
            font-size: 1.5rem;
        }}

        .pdf-name {{
            font-size: 0.9rem;
            font-weight: 500;
            color: var(--text-primary);
            word-break: break-all;
        }}

        .pdf-meta {{
            font-size: 0.75rem;
            color: var(--text-secondary);
            margin-top: 4px;
        }}

        /* Main Content Panel */
        .main-content {{
            flex: 1;
            display: flex;
            flex-direction: column;
            overflow-y: auto;
            background: radial-gradient(circle at top right, #11182c, #0b0f19);
        }}

        .header {{
            padding: 24px 32px;
            border-bottom: 1px solid var(--border-color);
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}

        .header-title h1 {{
            font-size: 1.75rem;
            font-weight: 700;
            letter-spacing: -0.5px;
            background: linear-gradient(to right, #ffffff, #93c5fd);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}

        .header-title p {{
            font-size: 0.9rem;
            color: var(--text-secondary);
            margin-top: 4px;
        }}

        .dashboard-grid {{
            padding: 32px;
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 24px;
            max-width: 1400px;
            margin: 0 auto;
            width: 100%;
        }}

        /* Card stylings */
        .card {{
            background-color: var(--card-bg);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 24px;
            box-shadow: var(--shadow);
            backdrop-filter: blur(10px);
            transition: transform 0.2s ease;
        }}

        .card:hover {{
            transform: translateY(-2px);
        }}

        .card-title {{
            font-size: 1.1rem;
            font-weight: 600;
            margin-bottom: 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            color: #ffffff;
            border-bottom: 1px solid rgba(255,255,255,0.06);
            padding-bottom: 12px;
        }}

        /* Status indicator */
        .status-badge {{
            background-color: rgba(16, 185, 129, 0.1);
            border: 1px solid var(--accent-green);
            color: var(--accent-green);
            padding: 6px 12px;
            border-radius: 50px;
            font-size: 0.8rem;
            font-weight: 500;
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }}

        .status-pulse {{
            width: 8px;
            height: 8px;
            background-color: var(--accent-green);
            border-radius: 50%;
            animation: pulse 1.5s infinite;
        }}

        @keyframes pulse {{
            0% {{ transform: scale(0.95); box-shadow: 0 0 0 0 rgba(16, 185, 129, 0.7); }}
            70% {{ transform: scale(1); box-shadow: 0 0 0 8px rgba(16, 185, 129, 0); }}
            100% {{ transform: scale(0.95); box-shadow: 0 0 0 0 rgba(16, 185, 129, 0); }}
        }}

        /* Timing Table */
        .timing-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 0.9rem;
        }}

        .timing-table th {{
            text-align: left;
            padding: 12px;
            color: var(--text-secondary);
            font-weight: 500;
            border-bottom: 1px solid var(--border-color);
        }}

        .timing-table td {{
            padding: 14px 12px;
            border-bottom: 1px solid rgba(255,255,255,0.04);
        }}

        .step-col {{
            color: var(--accent-blue);
            font-weight: 600;
            width: 50px;
        }}

        .checkpoint-col {{
            font-weight: 500;
            color: #ffffff;
        }}

        .duration-col {{
            font-family: 'JetBrains Mono', monospace;
            color: var(--accent-green);
            font-weight: 600;
            width: 100px;
        }}

        .desc-col {{
            color: var(--text-secondary);
            font-size: 0.85rem;
        }}

        /* Chart items */
        .chart-list {{
            display: flex;
            flex-direction: column;
            gap: 16px;
        }}

        .chart-item {{
            display: flex;
            flex-direction: column;
            gap: 6px;
        }}

        .chart-label {{
            display: flex;
            justify-content: space-between;
            font-size: 0.85rem;
            color: var(--text-secondary);
        }}

        .chart-bar-container {{
            height: 10px;
            background-color: rgba(255,255,255,0.06);
            border-radius: 10px;
            overflow: hidden;
        }}

        .chart-bar {{
            height: 100%;
            border-radius: 10px;
            transition: width 1s ease-in-out;
        }}

        /* Markdown / JSON View */
        .output-panel {{
            grid-column: span 2;
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 24px;
        }}

        .code-container {{
            background-color: #090d16;
            border-radius: 10px;
            padding: 16px;
            height: 350px;
            overflow: auto;
            border: 1px solid rgba(255,255,255,0.05);
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85rem;
            line-height: 1.5;
            color: #bbf7d0;
            white-space: pre-wrap;
        }}

        .markdown-container {{
            background-color: #ffffff;
            color: #1f2937;
            border-radius: 10px;
            padding: 20px;
            height: 350px;
            overflow: auto;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
            font-size: 0.95rem;
            line-height: 1.6;
        }}

        .markdown-container h1, .markdown-container h2, .markdown-container h3 {{
            margin-top: 16px;
            margin-bottom: 8px;
            font-weight: 600;
            border-bottom: 1px solid #e5e7eb;
            padding-bottom: 4px;
        }}

        .markdown-container p {{
            margin-bottom: 12px;
        }}

        .markdown-container pre {{
            background-color: #f3f4f6;
            padding: 12px;
            border-radius: 6px;
            overflow: auto;
            margin-bottom: 12px;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85rem;
        }}

        .markdown-container table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 16px;
        }}

        .markdown-container th, .markdown-container td {{
            border: 1px solid #d1d5db;
            padding: 8px 12px;
            text-align: left;
        }}

        .markdown-container th {{
            background-color: #f3f4f6;
        }}

        .total-time-card {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            background: linear-gradient(135deg, rgba(59, 130, 246, 0.1), rgba(139, 92, 246, 0.1));
            border-color: rgba(59, 130, 246, 0.3);
        }}

        .total-time-val {{
            font-size: 2rem;
            font-weight: 700;
            color: var(--accent-blue);
            font-family: 'JetBrains Mono', monospace;
        }}
    </style>
</head>
<body>

    <!-- Sidebar -->
    <div class="sidebar">
        <div class="sidebar-header">
            <div class="sidebar-title">
                <span>⚡</span> GLM-OCR Input PDFs
            </div>
        </div>
        <div class="sidebar-list">
            {sidebar_items}
        </div>
    </div>

    <!-- Main Content -->
    <div class="main-content">
        <div class="header">
            <div class="header-title">
                <h1>⚡ GLM-OCR Runtime Dashboard</h1>
                <p>Real-time document parsing and structural extraction benchmark</p>
            </div>
            <div>
                <span class="status-badge">
                    <span class="status-pulse"></span>
                    MaaS API Active
                </span>
            </div>
        </div>

        <div class="dashboard-grid">
            <!-- Summary Info -->
            <div class="card total-time-card" style="grid-column: span 2;">
                <div>
                    <h3 style="font-weight: 500; font-size: 1rem; color: var(--text-secondary);">Benchmark Target: <strong style="color: #ffffff;">{benchmark_data["filename"]}</strong></h3>
                    <p style="font-size: 0.85rem; color: var(--text-secondary); margin-top: 4px;">Pages: {benchmark_data["pages"]} | File Size: {benchmark_data["size_mb"]:.2f} MB</p>
                </div>
                <div>
                    <span style="font-size: 0.85rem; color: var(--text-secondary); display: block; text-align: right; margin-bottom: 4px;">Total Execution Time</span>
                    <span class="total-time-val">{benchmark_data["total_time"]:.2f}s</span>
                </div>
            </div>

            <!-- Checkpoint table -->
            <div class="card">
                <div class="card-title">
                    ⏱️ Checkpoint Step Timings
                </div>
                <table class="timing-table">
                    <thead>
                        <tr>
                            <th>Step</th>
                            <th>Checkpoint</th>
                            <th>Duration</th>
                            <th>Details</th>
                        </tr>
                    </thead>
                    <tbody>
                        {table_rows}
                    </tbody>
                </table>
            </div>

            <!-- CSS Bar Chart visualizer -->
            <div class="card">
                <div class="card-title">
                    📊 Processing Duration Breakdown
                </div>
                <div class="chart-list">
                    {chart_bars}
                </div>
            </div>

            <!-- Outputs Grid -->
            <div class="output-panel">
                <!-- Markdown Output View -->
                <div class="card">
                    <div class="card-title">
                        📝 Markdown View (Saved to {Path(benchmark_data["md_path"]).name})
                    </div>
                    <div class="markdown-container">
                        {render_mock_markdown_html(benchmark_data["markdown"])}
                    </div>
                </div>

                <!-- JSON Output View -->
                <div class="card">
                    <div class="card-title">
                        ⚙️ JSON Layout Structure (Saved to {Path(benchmark_data["json_path"]).name})
                    </div>
                    <div class="code-container">{json_formatted}</div>
                </div>
            </div>
        </div>
    </div>

</body>
</html>
"""
    return html_content


def render_mock_markdown_html(md_text):
    """Simple parser to render Markdown text to clean HTML for display."""
    import re

    html = md_text
    # Code blocks
    html = re.sub(r"```python\n(.*?)\n```", r"<pre><code>\1</code></pre>", html, flags=re.DOTALL)
    # Math blocks
    html = re.sub(
        r"\$\$\n(.*?)\n\$\$",
        r'<pre style="text-align: center; background-color: #f0fdf4; color: #166534;"><code>\1</code></pre>',
        html,
        flags=re.DOTALL,
    )

    # Tables
    # Identify tables and wrap in <table> tags
    def table_repl(match):
        lines = match.group(0).strip().split("\n")
        table_html = "<table>"
        for idx, line in enumerate(lines):
            if "---|" in line or "--|" in line:
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            tag = "th" if idx == 0 else "td"
            table_html += "<tr>" + "".join(f"<{tag}>{cell}</{tag}>" for cell in cells) + "</tr>"
        table_html += "</table>"
        return table_html

    html = re.sub(r"(\|.*?\|\n)+", table_repl, html)
    # Headers
    html = re.sub(r"^### (.*?)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.*?)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.*?)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)

    # Linebreaks / paras
    html = html.replace("\n\n", "<br>")
    return html


def main():
    pdfs = scan_pdfs()
    if not pdfs:
        print("No PDFs found in apps/input directory.")
        return

    print("Available PDFs:")
    for idx, pdf in enumerate(pdfs):
        print(f" {idx + 1}. {pdf.name} ({pdf.stat().st_size / (1024 * 1024):.2f} MB)")

    # By default, process the smallest PDF for timing efficiency
    target_pdf = pdfs[0]

    # Allow passing file index or name as argument
    if len(sys.argv) > 1:
        arg = sys.argv[1]
        if arg.isdigit():
            idx = int(arg) - 1
            if 0 <= idx < len(pdfs):
                target_pdf = pdfs[idx]
        else:
            for pdf in pdfs:
                if arg.lower() in pdf.name.lower():
                    target_pdf = pdf
                    break

    benchmark_data = parse_pdf_with_benchmarks(target_pdf)
    if benchmark_data:
        # Generate Dashboard HTML
        dashboard_html = generate_dashboard_html(benchmark_data, pdfs)

        dashboard_file = OUTPUT_DIR / "dashboard.html"
        with open(dashboard_file, "w", encoding="utf-8") as f:
            f.write(dashboard_html)

        word_report_file = OUTPUT_DIR / f"{target_pdf.stem}_performance_benchmark.docx"
        exported_report = export_word_report(benchmark_data, word_report_file)

        print("\n==============================================")
        print(f"Dashboard generated: {dashboard_file}")
        if exported_report:
            print(f"Word report generated: {exported_report}")
        print("==============================================")
        print("Open this file in your browser to view the timing benchmark and layout results!")


if __name__ == "__main__":
    main()
